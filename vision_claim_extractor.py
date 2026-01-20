"""
vision_claim_extractor.py

One script. One GPT-5 Vision call.
Inputs : bill image + prescription image
Outputs: claim_template_payload.json + ECHS_Claim_filled.docx

Usage:
  python vision_claim_extractor.py bill.jpeg prescription.jpeg
  python vision_claim_extractor.py bill.jpeg prescription.jpeg ECHS_Claim_template.docx

Requirements:
  pip install openai python-docx
  OPENAI_API_KEY set in environment
"""

import base64
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from openai import OpenAI


# ================= CONFIG =================
MODEL = "gpt-5"
OUT_JSON = "claim_template_payload.json"
OUT_DOCX = "ECHS_Claim_filled.docx"

# ONLY these keys will be bold + underlined
BOLD_UNDERLINE_KEYS = {
    "PATIENT_NAME",
    "ECHS_CARD_NO",
    "SERVICE_NO",
    "TOTAL_AMOUNT",
    "AMOUNT_WORDS",
}
# =========================================


def b64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode("utf-8")


def extract_json(text: str) -> dict:
    text = re.sub(r"```.*?\n", "", text, flags=re.S)
    start, end = text.find("{"), text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("No JSON found in model output")
    return json.loads(text[start:end + 1])


def to_money_2dp(x) -> str:
    if not x:
        return ""
    s = str(x).replace("₹", "").replace(",", "")
    m = re.search(r"\d+(\.\d+)?", s)
    return f"{float(m.group()):.2f}" if m else ""


def to_rupees_slash(x) -> str:
    if not x:
        return ""
    s = str(x).replace("₹", "").replace("/-", "").replace(",", "")
    m = re.search(r"\d+(\.\d+)?", s)
    return f"₹ {int(float(m.group()))} /-" if m else ""


def extract_with_gpt5(bill_img: Path, rx_img: Path) -> dict:
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    prompt = """
Extract data from the bill image and ECHS prescription image.

Return ONE JSON object ONLY.
No markdown. No commentary.

Rules:
- Dates: DD-MM-YYYY
- TOTAL_WO_DISCOUNT = subtotal before discount (2 decimals, no ₹)
- TOTAL_AMOUNT = payable amount (2 decimals, no ₹)
- Medicines: only purchased items (max 5)

Required keys (all must exist):

{
  "PATIENT_NAME": "",
  "ECHS_CARD_NO": "",
  "SERVICE_NO": "",
  "MOBILE_NO": "",
  "DIAGNOSIS": "",
  "INVOICE_NO.": "",
  "DATE": "",
  "DATE_EXPENDITURE": "",
  "TOTAL_WO_DISCOUNT": "",
  "TOTAL_AMOUNT": "",
  "AMOUNT_WORDS": "",
  "MED_1": "", "FORM_MED_1": "", "QTY_MED_1": "", "AMT_1": "",
  "MED_2": "", "FORM_MED_2": "", "QTY_MED_2": "", "AMT_2": "",
  "MED_3": "", "FORM_MED_3": "", "QTY_MED_3": "", "AMT_3": "",
  "MED_4": "", "FORM_MED_4": "", "QTY_MED_4": "", "AMT_4": "",
  "MED_5": "", "FORM_MED_5": "", "QTY_MED_5": "", "AMT_5": ""
}
"""

    resp = client.responses.create(
        model=MODEL,
        # temperature=0.0,
        input=[{
            "role": "user",
            "content": [
                {"type": "input_text", "text": prompt},
                {"type": "input_image", "image_url": f"data:image/jpeg;base64,{b64(bill_img)}"},
                {"type": "input_image", "image_url": f"data:image/jpeg;base64,{b64(rx_img)}"},
            ],
        }],
    )

    out_text = resp.output_text
    if not out_text:
        raise RuntimeError("GPT-5 returned empty output_text")

    return extract_json(out_text)


def normalise(data: dict) -> dict:
    out = dict(data)

    # Dates
    if not out.get("DATE_EXPENDITURE"):
        out["DATE_EXPENDITURE"] = out.get("DATE", "")

    out["CURRENT_MONTH_YEAR"] = datetime.now().strftime("%b %Y")

    # Money formatting
    out["TOTAL_AMOUNT"] = to_rupees_slash(out.get("TOTAL_AMOUNT"))
    out["TOTAL_WO_DISCOUNT"] = f"₹ {to_money_2dp(out.get('TOTAL_WO_DISCOUNT'))}"

    for i in range(1, 6):
        out[f"AMT_{i}"] = to_money_2dp(out.get(f"AMT_{i}"))

    return out


def replace_para(p, data):
    for k, v in data.items():
        ph = f"{{{{{k}}}}}"
        if ph not in p.text:
            continue

        if p.text.strip() == ph and k in BOLD_UNDERLINE_KEYS:
            p.clear()
            r = p.add_run(v)
            r.bold = True
            r.underline = True
            r.font.color.rgb = None  # ← CRITICAL
        else:
            for r in p.runs:
                r.text = r.text.replace(ph, v)

from docx.shared import Inches

def fix_certified_statements(doc):
    mapping = {
        "(1)": "(a)",
        "(2)": "(b)",
    }

    for p in doc.paragraphs:
        text = p.text.strip()
        for old, new in mapping.items():
            if text.startswith(old):
                p.text = text.replace(old, new, 1)
                p.paragraph_format.left_indent = Inches(0.5)



def fill_docx(template: Path, out_path: Path, data: dict):
    doc = Document(template)

    for p in doc.paragraphs:
        replace_para(p, data)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    replace_para(p, data)

    fix_certified_statements(doc)
    doc.save(out_path)



def main():
    if len(sys.argv) not in (3, 4):
        print("Usage: python vision_claim_extractor.py <bill> <prescription> [template]")
        sys.exit(1)

    bill = Path(sys.argv[1])
    rx = Path(sys.argv[2])
    template = Path(sys.argv[3]) if len(sys.argv) == 4 else Path("ECHS_Claim_template.docx")

    data = extract_with_gpt5(bill, rx)
    data = normalise(data)

    Path(OUT_JSON).write_text(json.dumps(data, indent=2), encoding="utf-8")
    fill_docx(template, Path(OUT_DOCX), data)

    print("✅ Generated:")
    print(" -", OUT_JSON)
    print(" -", OUT_DOCX)


if __name__ == "__main__":
    main()
